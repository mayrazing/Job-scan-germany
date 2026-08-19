import logging
import threading
import time
import traceback
from concurrent.futures import ThreadPoolExecutor
from importlib import import_module
from multiprocessing import active_children
from pathlib import Path
from types import ModuleType

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

FIXTURES = Path(__file__).parent / "fixtures" / "resume"


def resume_module() -> ModuleType:
    return import_module("job_scan.resume")


def _logging_state() -> object:
    def logger_state(logger: logging.Logger) -> object:
        return (
            logger.level,
            tuple(logger.handlers),
            tuple(logger.filters),
            logger.disabled,
            logger.propagate,
        )

    named_loggers = {
        name: logger_state(logger)
        for name, logger in logging.root.manager.loggerDict.items()
        if isinstance(logger, logging.Logger)
    }
    return (
        logging.getLogRecordFactory(),
        logging.root.manager.disable,
        logger_state(logging.root),
        named_loggers,
    )


def _write_content_damaged_pdf(path: Path, private_token: str) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=100, height=100)

    unicode_map = DecodedStreamObject()
    unicode_map.set_data(
        b"1 beginbfchar\n<00> <"
        + private_token.encode("ascii")
        + b">\nendbfchar"
    )
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
            NameObject("/Encoding"): NameObject("/WinAnsiEncoding"),
            NameObject("/ToUnicode"): writer._add_object(unicode_map),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )

    content = DecodedStreamObject()
    content.set_data(
        b"BT /F1 12 Tf (" + private_token.encode("ascii")
    )
    page[NameObject("/Contents")] = writer._add_object(content)
    writer.write(path)


def _wait_without_receiving_pdf_input(connection: object) -> None:
    time.sleep(60)
    connection.close()


class _TrackingConnection:
    def __init__(self) -> None:
        self.closed = False

    def close(self) -> None:
        self.closed = True


class _StartFailureProcess:
    pid: int | None = None

    def __init__(self, start_error: BaseException) -> None:
        self._start_error = start_error
        self.closed = False
        self.terminated = False

    def start(self) -> None:
        if isinstance(self._start_error, KeyboardInterrupt):
            self.pid = 12345
        raise self._start_error

    def is_alive(self) -> bool:
        return self.pid is not None and not self.terminated

    def join(self, _timeout: float | None = None) -> None:
        return None

    def terminate(self) -> None:
        self.terminated = True

    def kill(self) -> None:
        self.terminated = True

    def close(self) -> None:
        self.closed = True


class _StartFailureContext:
    def __init__(self, start_error: BaseException) -> None:
        self.connections = (_TrackingConnection(), _TrackingConnection())
        self.process = _StartFailureProcess(start_error)

    def Pipe(self, *, duplex: bool) -> tuple[_TrackingConnection, _TrackingConnection]:
        assert duplex
        return self.connections

    def Process(self, **_kwargs: object) -> _StartFailureProcess:
        return self.process


def _add_pdf_text_page(writer: PdfWriter, text: str) -> None:
    page = writer.add_blank_page(width=300, height=300)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )
    escaped_text = (
        text.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .replace("\r", "\\r")
        .replace("\n", "\\n")
    )
    content = DecodedStreamObject()
    content.set_data(f"BT /F1 12 Tf 20 250 Td ({escaped_text}) Tj ET".encode())
    page[NameObject("/Contents")] = writer._add_object(content)


@pytest.mark.parametrize("name", ["sample.pdf", "sample.docx"])
def test_extracts_text_resume(name: str) -> None:
    resume = resume_module()

    result = resume.extract_resume(FIXTURES / name)

    assert "Python" in result.text
    assert result.sha256.startswith("sha256:")
    assert result.path == FIXTURES / name
    assert result.format == name.removeprefix("sample.")


@pytest.mark.parametrize(
    ("name", "expected_hash"),
    [
        (
            "sample.pdf",
            "sha256:0568007cd98041124692550bf2d424d8d5ff89a0acb1f407d89cd042800d5bf9",
        ),
        (
            "sample.docx",
            "sha256:ce8d12508f4b064b099d92d890c66446807a7089e7b11c8efe7a0605cb4297ff",
        ),
    ],
)
def test_hashes_original_bytes_without_modifying_or_copying_fixture(
    name: str,
    expected_hash: str,
) -> None:
    resume = resume_module()
    path = FIXTURES / name
    original_bytes = path.read_bytes()
    original_entries = {child.name for child in FIXTURES.iterdir()}

    result = resume.extract_resume(path)

    assert result.sha256 == expected_hash
    assert path.read_bytes() == original_bytes
    assert {child.name for child in FIXTURES.iterdir()} == original_entries


def test_docx_keeps_paragraphs_and_table_cells_in_document_order(
    tmp_path: Path,
) -> None:
    resume = resume_module()
    path = tmp_path / "ordered.docx"
    document = Document()
    document.add_paragraph(
        "FIRST paragraph describes Python backend engineering experience."
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "SECOND skill"
    table.cell(0, 1).text = "THIRD evidence"
    table.cell(1, 0).text = "FOURTH testing"
    table.cell(1, 1).text = "FIFTH delivery"
    document.add_paragraph(
        "SIXTH paragraph adds enough real resume detail for validation."
    )
    document.save(path)

    result = resume.extract_resume(path)

    assert result.text == (
        "FIRST paragraph describes Python backend engineering experience.\n\n"
        "SECOND skill\n\n"
        "THIRD evidence\n\n"
        "FOURTH testing\n\n"
        "FIFTH delivery\n\n"
        "SIXTH paragraph adds enough real resume detail for validation."
    )


def test_pdf_keeps_non_empty_pages_in_order_and_normalizes_line_endings(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    resume = resume_module()
    path = tmp_path / "normalized.pdf"
    path.write_bytes(b"controlled PDF parser input")

    class Page:
        def __init__(self, text: str | None) -> None:
            self.text = text

        def extract_text(self) -> str | None:
            return self.text

    class Reader:
        def __init__(self, _source: object, **_kwargs: object) -> None:
            self.pages = [
                Page(
                    "FIRST Python page.\r\n\r\n\r\n"
                    "Backend engineering, APIs, SQL, and production delivery."
                ),
                Page(" \r\n "),
                Page(
                    "SECOND page.\rMore testing, observability, and team experience."
                ),
            ]

    monkeypatch.setattr(resume, "PdfReader", Reader)

    extracted = resume._extract_pdf_in_process(path.read_bytes())
    normalized = resume._normalize_text(extracted)

    assert normalized == (
        "FIRST Python page.\n\n"
        "Backend engineering, APIs, SQL, and production delivery.\n\n"
        "SECOND page.\nMore testing, observability, and team experience."
    )


def test_public_pdf_extraction_keeps_real_pages_in_order_and_normalizes_text(
    tmp_path: Path,
) -> None:
    resume = resume_module()
    path = tmp_path / "ordered.pdf"
    writer = PdfWriter()
    _add_pdf_text_page(
        writer,
        "FIRST Python page.\r\n\r\n\r\n"
        "Backend engineering, APIs, SQL, and production delivery.",
    )
    _add_pdf_text_page(
        writer,
        "SECOND page.\rMore testing, observability, and team experience.",
    )
    writer.write(path)

    result = resume.extract_resume(path)

    assert result.text == (
        "FIRST Python page.\n\n"
        "Backend engineering, APIs, SQL, and production delivery.\n\n"
        "SECOND page.\nMore testing, observability, and team experience."
    )


def test_pdf_input_send_uses_deadline_and_cleans_child_and_sender(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    resume = resume_module()
    monkeypatch.setattr(resume, "_PDF_EXTRACTION_TIMEOUT_SECONDS", 0.1)
    monkeypatch.setattr(resume, "_PDF_PROCESS_JOIN_SECONDS", 0.05)
    monkeypatch.setattr(
        resume,
        "_extract_pdf_worker",
        _wait_without_receiving_pdf_input,
    )
    completed = threading.Event()
    errors: list[BaseException] = []

    def extract() -> None:
        try:
            resume._extract_pdf(b"x" * (4 * 1024 * 1024))
        except BaseException as error:  # noqa: BLE001
            errors.append(error)
        finally:
            completed.set()

    caller = threading.Thread(target=extract, daemon=True)
    caller.start()
    completed_within_deadline = completed.wait(1.0)
    if not completed_within_deadline:
        for child in active_children():
            if child.name == resume._PDF_PROCESS_NAME:
                child.terminate()
                child.join(1.0)
    caller.join(1.0)

    assert completed_within_deadline
    assert errors and isinstance(errors[0], RuntimeError)
    assert all(child.name != resume._PDF_PROCESS_NAME for child in active_children())
    assert all(
        thread.name != "job-scan-pdf-sender" for thread in threading.enumerate()
    )


@pytest.mark.parametrize(
    "start_error",
    [OSError("controlled start failure"), KeyboardInterrupt()],
)
def test_pdf_process_is_closed_when_start_fails_or_parent_is_interrupted(
    monkeypatch: pytest.MonkeyPatch,
    start_error: BaseException,
) -> None:
    resume = resume_module()
    context = _StartFailureContext(start_error)
    monkeypatch.setattr(resume, "get_context", lambda _method: context)

    if isinstance(start_error, KeyboardInterrupt):
        with pytest.raises(KeyboardInterrupt):
            resume._extract_pdf(b"controlled input")
    else:
        with pytest.raises(RuntimeError, match="PDF extraction failed"):
            resume._extract_pdf(b"controlled input")

    assert context.process.closed
    assert all(connection.closed for connection in context.connections)
    if isinstance(start_error, KeyboardInterrupt):
        assert context.process.terminated


@pytest.mark.parametrize(
    ("fixture_name", "uppercase_suffix"),
    [("sample.pdf", ".PDF"), ("sample.docx", ".DOCX")],
)
def test_accepts_supported_suffixes_case_insensitively(
    tmp_path: Path,
    fixture_name: str,
    uppercase_suffix: str,
) -> None:
    resume = resume_module()
    path = tmp_path / f"resume{uppercase_suffix}"
    path.write_bytes((FIXTURES / fixture_name).read_bytes())

    result = resume.extract_resume(path)

    assert "Python" in result.text
    assert result.format == uppercase_suffix[1:].lower()


@pytest.mark.parametrize("suffix", [".txt", ".odt", ""])
def test_rejects_unsupported_suffix_with_actionable_domain_error(
    tmp_path: Path,
    suffix: str,
) -> None:
    resume = resume_module()
    secret = "PRIVATE resume content must not appear in errors"
    path = tmp_path / f"resume{suffix}"
    path.write_text(secret, encoding="utf-8")

    with pytest.raises(resume.UnsupportedResumeFormat) as caught:
        resume.extract_resume(path)

    message = str(caught.value)
    assert ".pdf" in message
    assert ".docx" in message
    assert secret not in message


def test_missing_resume_raises_domain_read_error(tmp_path: Path) -> None:
    resume = resume_module()

    with pytest.raises(resume.ResumeReadError, match="does not exist"):
        resume.extract_resume(tmp_path / "missing.pdf")


@pytest.mark.parametrize("suffix", [".pdf", ".docx"])
def test_malformed_supported_document_raises_safe_domain_error(
    tmp_path: Path,
    suffix: str,
) -> None:
    resume = resume_module()
    secret = "PRIVATE malformed resume contents"
    path = tmp_path / f"malformed{suffix}"
    path.write_text(secret, encoding="utf-8")

    with pytest.raises(resume.ResumeReadError) as caught:
        resume.extract_resume(path)

    assert secret not in str(caught.value)


def test_malformed_pdf_never_exposes_original_bytes_or_parser_error(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
    capsys: pytest.CaptureFixture[str],
) -> None:
    resume = resume_module()
    path = tmp_path / "malformed.pdf"
    path.write_bytes(b"PRIVATE malformed resume contents")
    caplog.set_level("DEBUG")

    with pytest.raises(resume.ResumeReadError) as caught:
        resume.extract_resume(path)

    captured = capsys.readouterr()
    rendered_exception = "".join(traceback.format_exception(caught.value))
    visible_channels = {
        "logs": "\n".join(record.getMessage() for record in caplog.records),
        "stdout": captured.out,
        "stderr": captured.err,
        "exception": str(caught.value),
        "exception traceback": rendered_exception,
    }
    for channel, visible_text in visible_channels.items():
        assert "PRIVA" not in visible_text, f"{channel} exposed the PDF header"
        assert "PRIVATE" not in visible_text, f"{channel} exposed resume bytes"
    assert "pypdf.errors." not in rendered_exception


def test_content_damaged_pdf_isolates_private_data_and_parser_failures(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
    capsys: pytest.CaptureFixture[str],
) -> None:
    resume = resume_module()
    private_token = "PRIVATE_STREAM_TOKEN_9A7B"
    path = tmp_path / "content-damaged.pdf"
    _write_content_damaged_pdf(path, private_token)
    caplog.set_level("DEBUG")
    logging.getLogger("pypdf._cmap")
    logger_state_before = _logging_state()

    def extract_and_catch() -> BaseException:
        try:
            resume.extract_resume(path)
        except BaseException as error:  # noqa: BLE001
            return error
        raise AssertionError("damaged PDF extraction unexpectedly succeeded")

    with ThreadPoolExecutor(max_workers=2) as pool:
        errors = list(pool.map(lambda _index: extract_and_catch(), range(2)))

    captured = capsys.readouterr()
    assert _logging_state() == logger_state_before
    assert all(child.name != resume._PDF_PROCESS_NAME for child in active_children())
    visible_channels = {
        "logs": "\n".join(record.getMessage() for record in caplog.records),
        "stdout": captured.out,
        "stderr": captured.err,
    }
    forbidden_fragments = (
        private_token,
        "PdfStreamError",
        "Stream has ended unexpectedly",
        "pypdf.errors.",
    )
    for index, error in enumerate(errors):
        assert isinstance(error, resume.ResumeReadError)
        assert error.__cause__ is None
        assert error.__context__ is None
        visible_channels[f"exception {index}"] = str(error)
        visible_channels[f"exception traceback {index}"] = "".join(
            traceback.format_exception(error)
        )
    for channel, visible_text in visible_channels.items():
        for fragment in forbidden_fragments:
            assert fragment not in visible_text, f"{channel} exposed {fragment}"


@pytest.mark.parametrize(
    ("character_count", "accepted"),
    [(99, False), (100, True)],
)
def test_pdf_enforces_exact_non_whitespace_character_boundary(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    character_count: int,
    accepted: bool,
) -> None:
    resume = resume_module()
    path = tmp_path / "boundary.pdf"
    path.write_bytes(b"controlled PDF parser input")

    class Page:
        def extract_text(self) -> str:
            return "x" * character_count

    class Reader:
        def __init__(self, _source: object, **_kwargs: object) -> None:
            self.pages = [Page()]

    monkeypatch.setattr(
        resume,
        "_extract_pdf",
        lambda _raw: Page().extract_text(),
    )

    if accepted:
        assert len(resume.extract_resume(path).text) == 100
    else:
        with pytest.raises(resume.ResumeTextMissing):
            resume.extract_resume(path)


def test_rejects_short_docx_with_domain_text_missing_error(tmp_path: Path) -> None:
    resume = resume_module()
    path = tmp_path / "short.docx"
    document = Document()
    document.add_paragraph("short text")
    document.save(path)

    with pytest.raises(resume.ResumeTextMissing, match="DOCX"):
        resume.extract_resume(path)


def test_rejects_scanned_or_empty_pdf_with_ocr_guidance(tmp_path: Path) -> None:
    resume = resume_module()
    path = tmp_path / "image-only.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.write(path)

    with pytest.raises(
        resume.ResumeTextMissing,
        match="OCR is not supported",
    ):
        resume.extract_resume(path)
